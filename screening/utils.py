"""
Helpers for pulling plain text out of uploaded resume files.

Supports:
  - .pdf   (via pypdf)
  - .docx  (via python-docx)
  - .txt   (plain decode, kept for backwards compatibility)
  - .zip   (a folder of any of the above, extracted and processed in bulk)
"""
import zipfile
from io import BytesIO

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}


class UnsupportedFileType(Exception):
    """Raised when a file's extension isn't one we know how to read."""


def extract_text_from_upload(uploaded_file):
    """
    Given a Django UploadedFile, return its extracted plain text.

    Raises UnsupportedFileType if the extension isn't .pdf, .docx, or .txt.
    Raises ValueError if the file can't be parsed (e.g. corrupt/encrypted PDF).
    """
    name = uploaded_file.name or ""
    return extract_text_from_bytes(name, uploaded_file.read())


def extract_text_from_bytes(filename, data):
    """
    Core extractor: given a filename (for its extension) and raw bytes,
    return the extracted plain text.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf_text(filename, data)
    elif ext == "docx":
        return _extract_docx_text(filename, data)
    elif ext == "txt":
        return data.decode("utf-8", errors="ignore")
    else:
        raise UnsupportedFileType(
            f"Unsupported file type '.{ext}'. Please upload a PDF or Word (.docx) file."
        )


def extract_resumes_from_zip(uploaded_zip_file):
    """
    Given a Django UploadedFile that is a .zip archive, extract every
    supported resume file inside it (recursively through folders).

    Returns a tuple: (candidates, skipped)
      - candidates: list of (candidate_name, resume_text) tuples
      - skipped: list of (filename, reason) tuples for files that couldn't be used
    """
    candidates = []
    skipped = []

    try:
        zf = zipfile.ZipFile(BytesIO(uploaded_zip_file.read()))
    except zipfile.BadZipFile:
        raise ValueError(f"'{uploaded_zip_file.name}' isn't a valid ZIP file.")

    for info in zf.infolist():
        # Skip directories and junk entries (macOS metadata, hidden files)
        if info.is_dir():
            continue
        filename = info.filename
        base = filename.rsplit("/", 1)[-1]
        if not base or base.startswith(".") or "__MACOSX" in filename:
            continue

        ext = base.rsplit(".", 1)[-1].lower() if "." in base else ""
        if ext not in SUPPORTED_EXTENSIONS:
            skipped.append((base, f"unsupported file type '.{ext}'"))
            continue

        try:
            data = zf.read(info)
            text = extract_text_from_bytes(base, data)
        except UnsupportedFileType as exc:
            skipped.append((base, str(exc)))
            continue
        except (ValueError, ImportError) as exc:
            skipped.append((base, str(exc)))
            continue

        candidate_name = base.rsplit(".", 1)[0]
        if text.strip():
            candidates.append((candidate_name, text))
        else:
            skipped.append((base, "no extractable text"))

    return candidates, skipped


def _extract_pdf_text(filename, data):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "PDF support requires the 'pypdf' package. Install it with: pip install pypdf"
        ) from exc

    try:
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError(f"'{filename}' is password-protected and can't be read.")

        pages_text = []
        for page in reader.pages:
            pages_text.append(page.extract_text() or "")
        return "\n".join(pages_text).strip()
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Couldn't read '{filename}' as a PDF: {exc}")


def _extract_docx_text(filename, data):
    try:
        import docx
    except ImportError as exc:
        raise ImportError(
            "Word support requires the 'python-docx' package. Install it with: pip install python-docx"
        ) from exc

    try:
        document = docx.Document(BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs]

        # Also pull text out of any tables, since resumes sometimes use them
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ValueError(f"Couldn't read '{filename}' as a Word document: {exc}")
